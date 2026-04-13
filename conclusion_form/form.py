import tkinter as tk
from tkinter import messagebox, BooleanVar
from tkinter import ttk
from tkinter.ttk import Combobox
from tkcalendar import Calendar
import datetime
import xml.etree.ElementTree as ET
from docx import Document
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer
import tempfile
import shutil
import re
import xml.dom.minidom as minidom
import ctypes
from ctypes import wintypes
import calendar

import sys, os, json

def setup_logging():
    log_dir = os.getcwd()
    log_file = os.path.join(log_dir, "log.txt")

    # Чтобы старые логи не затирались, можно дописывать
    sys.stdout = open(log_file, "a", encoding="utf-8")
    sys.stderr = sys.stdout

    print("\n=== Запуск программы:", datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "===\n")

def settings_path():
    path = resource_path("res/settings.json")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    return path

def resource_path(rel_path: str) -> str:
    """
    Возвращает путь к ресурсу и в dev-режиме, и внутри PyInstaller.
    rel_path: относительный путь внутри проекта (например 'conclusion_form/res/template.docx')
    """
    if hasattr(sys, '_MEIPASS'):
        base = sys._MEIPASS  # временная папка PyInstaller
    else:
        base = os.path.abspath(".")
    return os.path.join(base, rel_path)

SETTINGS_PATH = settings_path()
TEMPLATE_PATH = resource_path("conclusion_form/res/template.docx")
XML_PATH      = resource_path("conclusion_form/res/data.xml")
CALENDAR_PNG  = resource_path("conclusion_form/res/calendar.png")
PRIKAZ_XLSX   = resource_path("search_form/input/prikaz29n.xlsx")

class ConclusionForm(tk.Frame):
    def append_doc_with_page_break(self, master_path, appended_path):
        master_doc = Document(master_path)
        if master_doc.paragraphs:
            master_doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
        else:
            master_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        composer = Composer(master_doc)
        composer.append(Document(appended_path))
        composer.save(master_path)
    def __init__(self, parent, main_app):
        super().__init__(parent)
        self.main_app = main_app
        self.settings = main_app.settings
        self.suggestion_listbox = None
        self.report_org_window = None
        self.report_month_window = None
        self.records_editor_window = None
        self.records_selected_keys = set()
        self.records_all_visible_selected = False
        self.records_column_filters = {}
        self.records_cache = None
        self.records_search_after_id = None
        self.records_filter_windows = []

        # --- Переменные формы ---
        self.type_var = tk.StringVar(value="предварительный")
        self.organization = tk.StringVar()
        self.sex_var = tk.StringVar(value="М")
        self.division = tk.StringVar()
        self.profession = tk.StringVar()
        self.factors = tk.StringVar()
        self.typework = tk.StringVar()
        self.diagnosis = tk.StringVar()
        self.card_number = tk.StringVar()
        self.combine_all = BooleanVar(value=False)

        # Данные из XML
        self.data = self.load_data()

        # --- UI ---
        self.build_ui()

    # --------- UI строим здесь -------------
    def build_ui(self):
        self.pack(fill="both", expand=True)
        self.columnconfigure(0, weight=1)

        row = 0
        tk.Label(self, text="Дата ИДС").grid(row=row, column=0, sticky="w", padx=10, pady=(10, 0))
        self.ids_frame = tk.Frame(self)
        self.ids_frame.grid(row=row, column=1, sticky="w", padx=10)
        self.ids_entry = tk.Entry(self.ids_frame, width=45)
        self.ids_entry.pack(side="left", fill="x", expand=True)
        self.ids_entry.bind("<KeyRelease>", self.format_date)
        self.calendar_icon = tk.PhotoImage(master=self, file=CALENDAR_PNG  )
        tk.Button(
            self.ids_frame,
            image=self.calendar_icon,
            command=lambda: self.open_calendar(self.ids_entry),
            bd=0,
            relief="flat"
        ).pack(side="left", padx=(5, 0))

        row += 1
        tk.Label(self, text="Тип осмотра").grid(row=row, column=0, sticky="w", padx=10, pady=(10, 0))
        self.type_cb = Combobox(self, textvariable=self.type_var, values=["предварительный", "периодический"], width=50, state="readonly")
        self.type_cb.grid(row=row, column=1, padx=10, pady=(10, 0))

        row += 1
        tk.Label(self, text="Организация").grid(row=row, column=0, sticky="w", padx=10, pady=(10, 0))
        self.organization_cb = Combobox(self, textvariable=self.organization, width=50)
        self.organization_cb.grid(row=row, column=1, padx=10)
        self.organization_cb.bind("<<ComboboxSelected>>", self.on_organization_selected)

        row += 1
        tk.Label(self, text="ФИО").grid(row=row, column=0, sticky="w", padx=10, pady=(10, 0))
        self.name_entry = tk.Entry(self, width=53)
        self.name_entry.grid(row=row, column=1, padx=10)
        self.name_entry.bind("<FocusOut>", lambda e: self.sex_var.set(self.detect_sex_from_name(self.name_entry.get())))
        self.name_entry.bind("<KeyRelease>", self.show_name_suggestions)

        row += 1
        tk.Label(self, text="Дата рождения").grid(row=row, column=0, sticky="w", padx=10, pady=(10, 0))
        self.birthday_frame = tk.Frame(self)
        self.birthday_frame.grid(row=row, column=1, sticky="w", padx=10)
        self.birthday_entry = tk.Entry(self.birthday_frame, width=45)
        self.birthday_entry.pack(side="left", fill="x", expand=True)
        self.birthday_entry.bind("<KeyRelease>", self.format_date)
        tk.Button(
            self.birthday_frame,
            image=self.calendar_icon,
            command=lambda: self.open_calendar(self.birthday_entry),
            bd=0,
            relief="flat"
        ).pack(side="left", padx=(5, 0))

        row += 1
        tk.Label(self, text="Пол").grid(row=row, column=0, sticky="w", padx=10, pady=(10, 0))
        self.sex_cb = Combobox(self, textvariable=self.sex_var, values=["М", "Ж"], width=50, state="readonly")
        self.sex_cb.grid(row=row, column=1, padx=10)

        row += 1
        tk.Label(self, text="Подразделение").grid(row=row, column=0, sticky="w", padx=10, pady=(10, 0))
        self.division_cb = Combobox(self, textvariable=self.division, width=50)
        self.division_cb.grid(row=row, column=1, padx=10)

        row += 1
        tk.Label(self, text="Должность").grid(row=row, column=0, sticky="w", padx=10, pady=(10, 0))
        self.profession_cb = Combobox(self, textvariable=self.profession, width=50)
        self.profession_cb.grid(row=row, column=1, padx=10)

        row += 1
        tk.Label(self, text="Факторы").grid(row=row, column=0, sticky="w", padx=10, pady=(10, 0))
        self.factors_cb = Combobox(self, textvariable=self.factors, width=50)
        self.factors_cb.grid(row=row, column=1, padx=10)

        row += 1
        tk.Label(self, text="Виды работ").grid(row=row, column=0, sticky="w", padx=10, pady=(10, 0))
        self.typework_cb = Combobox(self, textvariable=self.typework, width=50)
        self.typework_cb.grid(row=row, column=1, padx=10)

        row += 1
        tk.Label(self, text="Диагноз").grid(row=row, column=0, sticky="w", padx=10, pady=(10, 0))
        self.diagnosis_cb = Combobox(self, textvariable=self.diagnosis, width=50)
        self.diagnosis_cb.grid(row=row, column=1, padx=10)
        self.diagnosis_cb.bind('<KeyRelease>', self.on_keyrelease)

        row += 1
        tk.Label(self, text="№ карты").grid(row=row, column=0, sticky="w", padx=10, pady=(10, 0))
        self.card_number_entry = tk.Entry(self, textvariable=self.card_number, width=53)
        self.card_number_entry.grid(row=row, column=1, padx=10)

        row += 1
        self.create_btn = tk.Button(self, text="Создать документ", command=self.generate_document, bg="#4CAF50", fg="white", height=2)
        self.create_btn.grid(row=row, column=0, columnspan=2, padx=10, pady=20, sticky="ew")

        row += 1
        self.combine_all = tk.BooleanVar(value=True)
        tk.Checkbutton(
            self,
            text="Объединить все в один файл",
            variable=self.combine_all
        ).grid(row=row, column=0, columnspan=2, pady=(0, 10), padx=10, sticky="w")

        # Уведомления (скрытый лейбл)
        row += 1
        self.notification_label = tk.Label(
            self,
            text="",
            fg="white",
            bg="#333",
            bd=1,
            relief="solid",
            padx=10, pady=5
        )
        self.notification_label.place_forget()

        self.edit_records_btn = tk.Button(
            self,
            text="Редактировать записи",
            command=self.open_records_editor,
            bg="#f2f2f2",
            fg="black"
        )
        self.edit_records_btn.place(relx=1.0, rely=1.0, anchor="se", x=-10, y=-10)

        # Автозаполнение
        self.organization_cb['values'] = sorted(self.data.keys())
        self.organization_cb.all_values = list(self.organization_cb['values'])

        for cb in (self.organization_cb, self.division_cb, self.profession_cb, self.factors_cb, self.typework_cb):
            cb.bind('<KeyRelease>', self.on_keyrelease)

        setup_logging()

    # ---------------- Логика ---------------------
    @staticmethod
    def sanitize_filename(name: str) -> str:
        return re.sub(r'[\\\/\:\*\?"<>\|]', '_', name)

    def show_name_suggestions(self, event):
        if self.suggestion_listbox:
            self.suggestion_listbox.destroy()
            self.suggestion_listbox = None

        text = self.name_entry.get().strip().lower()
        if not text:
            return

        all_names = []
        for records in self.data.values():
            for rec in records:
                fio = rec.get("name", "")
                if fio and fio.lower().startswith(text):
                    all_names.append(fio)
        suggestions = sorted(set(all_names))[:10]
        if not suggestions:
            return

        x = self.name_entry.winfo_rootx()
        y = self.name_entry.winfo_rooty() + self.name_entry.winfo_height()
        w = self.name_entry.winfo_width()
        h = min(200, len(suggestions) * 20)

        self.suggestion_listbox = tk.Toplevel(self)
        self.suggestion_listbox.overrideredirect(True)
        self.suggestion_listbox.transient(self)
        self.suggestion_listbox.geometry(f"{w}x{h}+{x}+{y}")
        self.suggestion_listbox.lift()

        lb = tk.Listbox(self.suggestion_listbox, exportselection=False)
        lb.pack(fill="both", expand=True)
        for item in suggestions:
            lb.insert(tk.END, item)

        def on_select(evt):
            sel = lb.get(lb.curselection())
            self.suggestion_listbox.destroy()
            self.suggestion_listbox = None
            self.fill_person_fields(sel)
            self.focus_force()
        lb.bind("<ButtonRelease-1>", on_select)

    def fill_person_fields(self, fio):
        for records in self.data.values():
            for rec in records:
                if rec.get("name") == fio:
                    self.type_var.set(rec.get("type", "предварительный"))
                    self.name_entry.delete(0, tk.END)
                    self.name_entry.insert(0, rec["name"])
                    self.birthday_entry.delete(0, tk.END)
                    self.birthday_entry.insert(0, rec["birthday"])
                    self.sex_var.set(rec["sex"])
                    self.card_number.set(rec.get("card_number", ""))
                    return

    @staticmethod
    def detect_sex_from_name(full_name):
        parts = full_name.strip().split()
        if len(parts) >= 3:
            middle = parts[2].lower()
            if middle.endswith(("вич", "льич", "ич")):
                return "М"
            elif middle.endswith(("вна", "чна", "инична", "овна", "евна", "ична")):
                return "Ж"
        return "М"

    def load_settings(self):
        if os.path.exists(SETTINGS_PATH):
            with open(SETTINGS_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        else:
            return {"save_dir": os.getcwd()}

    def save_settings(self, settings):
        with open(SETTINGS_PATH, "w", encoding="utf-8") as f:
            json.dump(settings, f, indent=2, ensure_ascii=False)

    def clear_form(self):
        self.type_var.set("предварительный")
        self.organization.set("")
        self.name_entry.delete(0, tk.END)
        self.birthday_entry.delete(0, tk.END)
        self.sex_var.set("М")
        self.division.set("")
        self.profession.set("")
        self.factors.set("")
        self.typework.set("")
        self.diagnosis.set("")
        self.card_number.set("")
        self.ids_entry.delete(0, tk.END)

    def load_data(self):
        if not os.path.exists(XML_PATH):
            return {}
        tree = ET.parse(XML_PATH)
        root = tree.getroot()
        data = {}
        for person in root.findall("person"):
            org_name = person.findtext("organization", default="")
            record = {
                "type": person.findtext("type", default="предварительный"),
                "name": person.findtext("name", default=""),
                "birthday": person.findtext("birthday", default=""),
                "sex": person.findtext("sex", default=""),
                "division": person.findtext("division", default=""),
                "profession": person.findtext("profession", default=""),
                "factors": person.findtext("factors", default=""),
                "typework": person.findtext("typework", default=""),
                "id": person.findtext("id", default=""),
                "ids_date": person.findtext("ids_date", default=""),
                "diagnosis": person.findtext("diagnosis", default=""),
                "card_number": person.findtext("card_number", default="")
            }
            if org_name not in data:
                data[org_name] = []
            data[org_name].append(record)
        return data

    def prettify_xml(self, xml_path):
        xml_str = open(xml_path, "r", encoding="utf-8").read()
        dom = minidom.parseString(xml_str)
        pretty_xml = dom.toprettyxml(indent="  ", encoding="utf-8")
        with open(xml_path, "wb") as f:
            f.write(pretty_xml)

    def save_record(self, org_name, division, profession, factors, typework,
                    name=None, birthday=None, sex_val=None, diagnosis=None, ids_date=None, card_number=None,
                    type_value=None):
        if not (name and birthday and sex_val):
            return
        if os.path.exists(XML_PATH):
            tree = ET.parse(XML_PATH)
            root = tree.getroot()
        else:
            root = ET.Element("data")
            tree = ET.ElementTree(root)
        now_str = datetime.datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        p = ET.SubElement(root, "person")
        ET.SubElement(p, "organization").text = org_name
        ET.SubElement(p, "type").text = type_value if type_value else "предварительный"
        ET.SubElement(p, "name").text = name
        ET.SubElement(p, "birthday").text = birthday
        ET.SubElement(p, "sex").text = sex_val
        ET.SubElement(p, "division").text = division
        ET.SubElement(p, "profession").text = profession
        ET.SubElement(p, "factors").text = factors
        ET.SubElement(p, "typework").text = typework
        ET.SubElement(p, "diagnosis").text = diagnosis if diagnosis else ""
        ET.SubElement(p, "card_number").text = card_number if card_number else ""
        ET.SubElement(p, "ids_date").text = ids_date if ids_date else ""
        ET.SubElement(p, "id").text = str(int(datetime.datetime.now().timestamp()))
        ET.SubElement(p, "date_created").text = now_str
        tree.write(XML_PATH, encoding="utf-8", xml_declaration=True)
        self.prettify_xml(XML_PATH)

    def get_unique_values(self, field, org_name=None):
        """Вернёт уникальные значения поля. Если задана org_name — только для этой организации."""
        values = set()

        if org_name and org_name in self.data:
            records = self.data[org_name]
        else:
            # все записи по всем организациям
            records = [rec for recs in self.data.values() for rec in recs]

        for record in records:
            val = record.get(field)
            if val:
                values.add(val)

        return sorted(values)

    def replace_placeholders(self, doc, data_dict):
        def replace_in_paragraph(paragraph, data_dict):
            text = ''.join(run.text for run in paragraph.runs)
            parts = re.split(r'(\{.*?\})', text)
            if not any(part in data_dict for part in parts):
                return
            paragraph.clear()
            for part in parts:
                if part in data_dict:
                    run = paragraph.add_run(data_dict[part])
                    run.bold = True
                else:
                    paragraph.add_run(part)

        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, data_dict)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, data_dict)

    def on_keyrelease(self, event):
        cb = event.widget
        txt = cb.get().lower()
        if not hasattr(cb, 'all_values'):
            cb.all_values = list(cb['values'])
        if txt == '':
            vals = cb.all_values
        else:
            vals = [v for v in cb.all_values if txt in v.lower()]
        cb['values'] = vals
        try:
            cb.tk.call('ttk::combobox::post', cb._w)
            cb.focus_force()
        except Exception:
            pass

    def generate_document(self):
        type_raw = self.type_var.get()
        type_genitive = {
            "предварительный": "предварительного",
            "периодический": "периодического"
        }.get(type_raw, type_raw)

        form_data = {
            "{type}": type_genitive,
            "{organization}": self.organization.get(),
            "{name}": self.name_entry.get(),
            "{birthday}": self.birthday_entry.get(),
            "{sex}": self.sex_var.get(),
            "{division}": self.division.get(),
            "{profession}": self.profession.get(),
            "{factors}": self.factors.get(),
            "{typework}": self.typework.get(),
            "{ids_date}": self.ids_entry.get(),
            "{diagnosis}": self.diagnosis.get(),
            "{card_number}": self.card_number.get(),
            "{year}": str(datetime.datetime.now().year)   # текущий год
        }

        if (not form_data["{organization}"] or
                not form_data["{name}"] or
                not form_data["{birthday}"] or
                not form_data["{ids_date}"]):
            messagebox.showerror(
                "Ошибка ввода",
                "Пожалуйста, заполните обязательные поля:\n"
                "• Организация\n"
                "• ФИО\n"
                "• Дата рождения\n"
                "• Дата ИДС"
            )
            return

        if not self.is_valid_date(self.birthday_entry.get()):
            messagebox.showerror(
                "Ошибка даты",
                "Дата рождения должна быть в формате ДД.ММ.ГГГГ"
            )
            return
        if self.ids_entry.get().strip() and not self.is_valid_date(self.ids_entry.get()):
            messagebox.showerror(
                "Ошибка даты",
                "Дата ИДС должна быть в формате ДД.ММ.ГГГГ"
            )
            return

        temp_doc = Document(TEMPLATE_PATH)
        self.replace_placeholders(temp_doc, form_data)

        temp_doc_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
        temp_doc.save(temp_doc_path)

        if self.combine_all.get():
            combined_filename = os.path.join(
                self.settings.get("save_dir", os.getcwd()),
                f"заключения_{datetime.datetime.now().strftime('%d.%m.%Y')}.docx"
            )

            if os.path.exists(combined_filename):
                try:
                    self.append_doc_with_page_break(combined_filename, temp_doc_path)
                    self.show_notification(f"Добавлено в файл: {combined_filename}")
                except Exception as e:
                    messagebox.showerror("Ошибка", f"Не удалось дописать в файл:\n{e}")
                    return
            else:
                # создаём новый файл из temp_doc
                try:
                    shutil.copyfile(temp_doc_path, combined_filename)
                    self.show_notification(f"Создан новый файл: {combined_filename}")
                    print("создан новый файл для дозаписи")
                except Exception as e:
                    messagebox.showerror("Ошибка", f"Не удалось создать файл:\n{e}")
                    print("файл не создан")
                    return

        else:
            filename = os.path.join(
                self.settings.get("save_dir", os.getcwd()),
                f"{form_data['{name}']} - заключение.docx"
            )
            try:
                temp_doc.save(filename)
                self.show_notification(f"Файл сохранён: {filename}")
            except PermissionError:
                messagebox.showerror(
                    "Ошибка записи",
                    "Невозможно создать запись, проверьте, закрыт ли Word-файл."
                )
                return
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\n{e}")
                return

        self.save_record(
            form_data["{organization}"],
            form_data["{division}"],
            form_data["{profession}"],
            form_data["{factors}"],
            form_data["{typework}"],
            name=form_data["{name}"],
            birthday=form_data["{birthday}"],
            sex_val=form_data["{sex}"],
            ids_date=self.ids_entry.get(),
            diagnosis=self.diagnosis.get(),
            card_number=self.card_number.get(),
            type_value=self.type_var.get()
        )
        self.data = self.load_data()
        self.organization_cb['values'] = sorted(self.data.keys())
        self.update_comboboxes()
        try:
            os.remove(temp_doc_path)
        except Exception as e:
            print(f"Не удалось удалить временный файл: {temp_doc_path}\n{e}")
        self.clear_form()

    def update_comboboxes(self):
        # всегда обновляем список организаций
        self.organization_cb["values"] = sorted(self.data.keys())

        org = self.organization.get().strip()
        if org and org in self.data:
            # значения ТОЛЬКО для выбранной организации
            self.division_cb["values"] = self.get_unique_values("division", org)
            self.profession_cb["values"] = self.get_unique_values("profession", org)
            self.factors_cb["values"] = self.get_unique_values("factors", org)
            self.typework_cb["values"] = self.get_unique_values("typework", org)
            self.diagnosis_cb["values"] = self.get_unique_values("diagnosis", org)
        else:
            # если организация не выбрана — показываем общие списки (как раньше)
            self.division_cb["values"] = self.get_unique_values("division")
            self.profession_cb["values"] = self.get_unique_values("profession")
            self.factors_cb["values"] = self.get_unique_values("factors")
            self.typework_cb["values"] = self.get_unique_values("typework")
            self.diagnosis_cb["values"] = self.get_unique_values("diagnosis")

        # обновляем кеш для живого поиска в combobox'ах
        for cb in (self.organization_cb, self.division_cb, self.profession_cb, self.factors_cb, self.typework_cb,
                   self.diagnosis_cb):
            cb.all_values = list(cb['values'])

    def on_organization_selected(self, event):
        self.update_comboboxes()

    @staticmethod
    def is_valid_date(date_str):
        try:
            datetime.datetime.strptime(date_str, "%d.%m.%Y")
            return True
        except ValueError:
            return False

    def format_date(self, event):
        widget = event.widget
        s = widget.get()
        digits = ''.join(filter(str.isdigit, s))[:8]
        parts = []
        if len(digits) >= 2:
            parts.append(digits[:2])
        else:
            parts.append(digits)
        if len(digits) >= 4:
            parts.append(digits[2:4])
        elif len(digits) > 2:
            parts.append(digits[2:])
        if len(digits) > 4:
            parts.append(digits[4:])
        new_text = '.'.join(parts)
        if new_text != s:
            widget.delete(0, tk.END)
            widget.insert(0, new_text)
            widget.icursor(tk.END)

    def open_calendar(self, entry_widget):
        mouse_x = self.winfo_pointerx()
        mouse_y = self.winfo_pointery()
        top = tk.Toplevel(self)
        top.overrideredirect(False)
        top.title("Выберите дату")
        def pick_date():
            date = cal.selection_get()
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, date.strftime("%d.%m.%Y"))
            top.destroy()
        cal = Calendar(top, date_pattern='dd.mm.yyyy')
        cal.pack(padx=10, pady=10)
        tk.Button(top, text="Выбрать", command=pick_date).pack(pady=5)
        top.update_idletasks()
        win_w = top.winfo_width()
        win_h = top.winfo_height()
        try:
            SPI_GETWORKAREA = 0x0030
            rect = wintypes.RECT()
            ctypes.windll.user32.SystemParametersInfoW(SPI_GETWORKAREA, 0, ctypes.byref(rect), 0)
            work_w = rect.right - rect.left
            work_h = rect.bottom - rect.top
        except Exception:
            work_w = self.winfo_screenwidth()
            work_h = self.winfo_screenheight()
        x = mouse_x + 10
        y = mouse_y + 10
        if x + win_w > work_w:
            x = work_w - win_w
        if y + win_h > work_h:
            y = work_h - win_h
        if x < 0:
            x = 0
        if y < 0:
            y = 0
        top.geometry(f"{win_w}x{win_h}+{x}+{y}")

    def show_notification(self, text, duration=3000, x_offset=10, y_offset=10):
        self.notification_label.config(text=text)
        self.notification_label.update_idletasks()
        self.notification_label.place(
            relx=1.0, rely=1.0,
            anchor="se",
            x=-x_offset,
            y=-y_offset
        )
        self.after(duration, self.notification_label.place_forget)

    # # ----------------- ОТЧЁТЫ --------------------
    # def report_by_organization(self):
    #     # -- Твой старый код "report_by_organization", только self везде --
    #     import pandas as pd
    #     if self.report_org_window and self.report_org_window.winfo_exists():
    #         self.report_org_window.focus_force()
    #         return
    #     self.report_org_window = tk.Toplevel(self)
    #     rpt = self.report_org_window
    #     rpt.title("Отчет по организации")
    #     rpt.resizable(False, False)
    #     rpt.protocol("WM_DELETE_WINDOW", lambda: (rpt.destroy(), self.set_report_none()))
    #     padx, pady = 10, 5
    #     org_var_report = tk.StringVar()
    #     tk.Label(rpt, text="Организация:").grid(row=0, column=0, sticky="w", padx=padx, pady=pady)
    #     org_list = sorted(self.data.keys())
    #     org_cb = Combobox(
    #         rpt,
    #         values=org_list,
    #         textvariable=org_var_report,
    #         width=40,
    #         state="readonly"
    #     )
    #     org_cb.grid(row=0, column=1, padx=padx, pady=pady)
    #     org_var_report.set("")
    #     tk.Label(rpt, text="Период с:").grid(row=1, column=0, sticky="w", padx=padx, pady=pady)
    #     start_var = tk.StringVar()
    #     start_entry = tk.Entry(rpt, width=20, textvariable=start_var)
    #     start_entry.grid(row=1, column=1, sticky="w", padx=padx, pady=pady)
    #     tk.Button(rpt, text="📅", command=lambda: self.open_calendar(start_entry)).grid(row=1, column=2, padx=0, pady=pady)
    #     tk.Label(rpt, text="По:").grid(row=2, column=0, sticky="w", padx=padx, pady=pady)
    #     end_var = tk.StringVar()
    #     end_entry = tk.Entry(rpt, width=20, textvariable=end_var)
    #     end_entry.grid(row=2, column=1, sticky="w", padx=padx, pady=pady)
    #     tk.Button(rpt, text="📅", command=lambda: self.open_calendar(end_entry)).grid(row=2, column=2, padx=0, pady=pady)
    #
    #     def on_start_changed(*_):
    #         s = start_var.get().strip()
    #         if self.is_valid_date(s):
    #             dt = datetime.datetime.strptime(s, "%d.%m.%Y")
    #             last = calendar.monthrange(dt.year, dt.month)[1]
    #             end_var.set(f"{last:02d}.{dt.month:02d}.{dt.year}")
    #
    #     start_var.trace_add("write", on_start_changed)
    #
    #     def make_report():
    #         org_sel = org_var_report.get().strip()
    #         start = start_entry.get().strip()
    #         end = end_entry.get().strip()
    #         if not org_sel:
    #             messagebox.showerror("Ошибка ввода", "Выберите организацию")
    #             return
    #         if not self.is_valid_date(start) or not self.is_valid_date(end):
    #             messagebox.showerror("Ошибка даты", "Даты в формате ДД.ММ.ГГГГ")
    #             return
    #         d0 = datetime.datetime.strptime(start, "%d.%m.%Y")
    #         d1 = datetime.datetime.strptime(end, "%d.%m.%Y")
    #         if d1 < d0:
    #             messagebox.showerror("Ошибка", "Конечная дата меньше начальной")
    #             return
    #         rows = []
    #         for r in self.data.get(org_sel, []):
    #             ids = r.get("ids_date", "").strip()
    #             if not ids:
    #                 continue
    #             try:
    #                 d_ids = datetime.datetime.strptime(ids, "%d.%m.%Y")
    #             except ValueError:
    #                 continue
    #             if d0 <= d_ids <= d1:
    #                 rows.append({
    #                     "Организация": org_sel,
    #                     "ФИО": r["name"],
    #                     "Дата рожд.": r["birthday"],
    #                     "Пол": r["sex"],
    #                     "Подразделение": r["division"],
    #                     "Должность": r["profession"],
    #                     "Факторы": r["factors"],
    #                     "Виды работ": r["typework"],
    #                     "Дата ИДС": ids,
    #                     "Диагноз": r.get("diagnosis", "")
    #                 })
    #         if not rows:
    #             messagebox.showinfo("Пустой отчет", "Нет записей за выбранный период.")
    #             return
    #         df = pd.DataFrame(rows)
    #         save_dir = self.settings.get("save_dir", os.getcwd())
    #         fname = self.sanitize_filename(f"Отчет_{org_sel}_{start}_{end}.xlsx")
    #         save_path = os.path.join(save_dir, fname)
    #         try:
    #             from openpyxl.utils import get_column_letter
    #             with pd.ExcelWriter(save_path, engine="openpyxl") as writer:
    #                 df.to_excel(writer, index=False, sheet_name="Report")
    #                 sheet = writer.sheets["Report"]
    #                 for idx, col in enumerate(df.columns, start=1):
    #                     width = max(df[col].astype(str).map(len).max(), len(col)) + 2
    #                     sheet.column_dimensions[get_column_letter(idx)].width = width
    #             self.show_notification(f"Отчет сохранён: {save_path}")
    #         except Exception as e:
    #             messagebox.showerror("Ошибка записи", f"Не удалось сохранить файл:\n{e}")
    #
    #     tk.Button(
    #         rpt,
    #         text="Сформировать",
    #         command=make_report,
    #         bg="#4CAF50",
    #         fg="white"
    #     ).grid(row=3, column=0, columnspan=3, pady=(10, 10), padx=padx, sticky="ew")
    #     rpt.grid_columnconfigure(1, weight=1)
    #     rpt.update_idletasks()
    #     w, h = rpt.winfo_width(), rpt.winfo_height()
    #     sw, sh = rpt.winfo_screenwidth(), rpt.winfo_screenheight()
    #     x, y = (sw - w) // 2, (sh - h) // 2
    #     rpt.geometry(f"{w}x{h}+{x}+{y}")
    #
    # def set_report_none(self):
    #     self.report_org_window = None
    #     self.report_month_window = None
    #
    # def report_by_month(self):
    #     import pandas as pd
    #     if self.report_month_window and self.report_month_window.winfo_exists():
    #         self.report_month_window.focus_force()
    #         return
    #     self.report_month_window = tk.Toplevel(self)
    #     rpt = self.report_month_window
    #     rpt.title("Отчет по дате ИДС")
    #     rpt.resizable(False, False)
    #     rpt.protocol("WM_DELETE_WINDOW", lambda: (rpt.destroy(), self.set_report_none()))
    #     padx, pady = 10, 5
    #     tk.Label(rpt, text="Период с:").grid(row=0, column=0, sticky="w", padx=padx, pady=pady)
    #     start_var = tk.StringVar()
    #     start_entry = tk.Entry(rpt, width=20, textvariable=start_var)
    #     start_entry.grid(row=0, column=1, sticky="w", padx=padx, pady=pady)
    #     tk.Button(rpt, text="📅", command=lambda: self.open_calendar(start_entry)).grid(row=0, column=2, padx=0, pady=pady)
    #     tk.Label(rpt, text="По:").grid(row=1, column=0, sticky="w", padx=padx, pady=pady)
    #     end_var = tk.StringVar()
    #     end_entry = tk.Entry(rpt, width=20, textvariable=end_var)
    #     end_entry.grid(row=1, column=1, sticky="w", padx=padx, pady=pady)
    #     tk.Button(rpt, text="📅", command=lambda: self.open_calendar(end_entry)).grid(row=1, column=2, padx=0, pady=pady)
    #     def on_start_changed(*_):
    #         s = start_var.get().strip()
    #         if self.is_valid_date(s):
    #             dt = datetime.datetime.strptime(s, "%d.%m.%Y")
    #             last = calendar.monthrange(dt.year, dt.month)[1]
    #             end_var.set(f"{last:02d}.{dt.month:02d}.{dt.year}")
    #     start_var.trace_add("write", on_start_changed)
    #     def make_report_month():
    #         start = start_var.get().strip()
    #         end = end_var.get().strip()
    #         d0 = datetime.datetime.strptime(start, "%d.%m.%Y")
    #         d1 = datetime.datetime.strptime(end, "%d.%m.%Y")
    #         rows = []
    #         for org_name, recs in self.data.items():
    #             for r in recs:
    #                 ids = r.get("ids_date", "").strip()
    #                 if not ids:
    #                     continue
    #                 try:
    #                     d_ids = datetime.datetime.strptime(ids, "%d.%m.%Y")
    #                 except ValueError:
    #                     continue
    #                 if d0 <= d_ids <= d1:
    #                     rows.append({
    #                         "Организация": org_name,
    #                         "ФИО": r["name"],
    #                         "Дата рожд.": r["birthday"],
    #                         "Пол": r["sex"],
    #                         "Подразделение": r["division"],
    #                         "Должность": r["profession"],
    #                         "Факторы": r["factors"],
    #                         "Виды работ": r["typework"],
    #                         "Дата ИДС": ids,
    #                         "Диагноз": r.get("diagnosis", "")
    #                     })
    #         if not rows:
    #             messagebox.showinfo("Пустой отчет", "Нет записей за выбранный период.")
    #             return
    #         df = pd.DataFrame(rows)
    #         save_dir = self.settings.get("save_dir", os.getcwd())
    #         fname = self.sanitize_filename(f"Отчет_по_месяцу_{start}_{end}.xlsx")
    #         save_path = os.path.join(save_dir, fname)
    #         try:
    #             from openpyxl.utils import get_column_letter
    #             with pd.ExcelWriter(save_path, engine="openpyxl") as writer:
    #                 df.to_excel(writer, index=False, sheet_name="Report")
    #                 sheet = writer.sheets["Report"]
    #                 for idx, col in enumerate(df.columns, start=1):
    #                     width = max(df[col].astype(str).map(len).max(), len(col)) + 2
    #                     sheet.column_dimensions[get_column_letter(idx)].width = width
    #             self.show_notification(f"Отчет сохранён: {save_path}")
    #         except Exception as e:
    #             messagebox.showerror("Ошибка записи", f"Не удалось сохранить файл:\n{e}")
    #     tk.Button(
    #         rpt,
    #         text="Сформировать",
    #         command=make_report_month,
    #         bg="#4CAF50",
    #         fg="white"
    #     ).grid(row=2, column=0, columnspan=3, pady=(10, 10), padx=padx, sticky="ew")
    #     rpt.grid_columnconfigure(1, weight=1)
    #     rpt.update_idletasks()
    #     w, h = rpt.winfo_width(), rpt.winfo_height()
    #     sw, sh = rpt.winfo_screenwidth(), rpt.winfo_screenheight()
    #     x, y = (sw - w) // 2, (sh - h) // 2
    #     rpt.geometry(f"{w}x{h}+{x}+{y}")

    def open_records_editor(self):
        if self.records_editor_window and self.records_editor_window.winfo_exists():
            self.records_editor_window.deiconify()
            self.records_editor_window.lift()
            self.records_editor_window.focus_force()
            return

        self.records_selected_keys = set()
        self.records_all_visible_selected = False
        self.records_cache = self._read_all_xml_records(force_reload=True)

        top = tk.Toplevel(self)
        self.records_editor_window = top
        top.title("Редактирование записей")
        top.geometry("1400x720")

        def on_close():
            self.close_records_filter_windows()
            if self.records_editor_window and self.records_editor_window.winfo_exists():
                self.records_editor_window.destroy()
            self.records_editor_window = None

        top.protocol("WM_DELETE_WINDOW", on_close)

        filters_frame = tk.Frame(top)
        filters_frame.pack(fill="x", padx=10, pady=10)

        tk.Label(filters_frame, text="Дата с:").grid(row=0, column=0, sticky="w")
        self.records_date_from_var = tk.StringVar()
        date_from_entry = tk.Entry(filters_frame, textvariable=self.records_date_from_var, width=18)
        date_from_entry.grid(row=0, column=1, padx=(6, 8), sticky="w")
        date_from_entry.bind("<KeyRelease>", self.format_date)
        tk.Button(filters_frame, text="📅", command=lambda: self.open_calendar(date_from_entry)).grid(row=0, column=2, padx=(0, 12))

        tk.Label(filters_frame, text="Дата по:").grid(row=0, column=3, sticky="w")
        self.records_date_to_var = tk.StringVar()
        date_to_entry = tk.Entry(filters_frame, textvariable=self.records_date_to_var, width=18)
        date_to_entry.grid(row=0, column=4, padx=(6, 8), sticky="w")
        date_to_entry.bind("<KeyRelease>", self.format_date)
        tk.Button(filters_frame, text="📅", command=lambda: self.open_calendar(date_to_entry)).grid(row=0, column=5, padx=(0, 12))

        tk.Label(filters_frame, text="Поиск по ФИО:").grid(row=0, column=6, sticky="w")
        self.records_search_var = tk.StringVar()
        search_entry = tk.Entry(filters_frame, textvariable=self.records_search_var, width=35)
        search_entry.grid(row=0, column=7, padx=(6, 8), sticky="ew")
        filters_frame.grid_columnconfigure(7, weight=1)

        tk.Button(filters_frame, text="Применить", command=self.apply_records_filters, bg="#4CAF50", fg="white").grid(row=0, column=8, padx=(4, 0))
        tk.Button(filters_frame, text="Сбросить", command=self.reset_records_filters).grid(row=0, column=9, padx=(6, 0))

        tree_frame = tk.Frame(top)
        tree_frame.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        columns = (
            "selected", "type", "organization", "name", "birthday", "sex", "division", "profession",
            "factors", "typework", "diagnosis", "card_number", "ids_date", "date_created"
        )

        self.records_tree = ttk.Treeview(tree_frame, columns=columns, show="headings")

        y_scroll = ttk.Scrollbar(tree_frame, orient="vertical", command=self.records_tree.yview)
        x_scroll = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.records_tree.xview)

        self.records_tree.grid(row=0, column=0, sticky="nsew")
        y_scroll.grid(row=0, column=1, sticky="ns")
        x_scroll.grid(row=1, column=0, sticky="ew")

        tree_frame.rowconfigure(0, weight=1)
        tree_frame.columnconfigure(0, weight=1)

        self.records_tree.configure(yscrollcommand=y_scroll.set, xscrollcommand=x_scroll.set)

        self.records_tree.heading("selected", text="☐", command=self.toggle_all_records_selection)
        self.records_tree.heading("type", text="Тип осмотра ▼", command=lambda: self.open_records_column_filter("type", "Тип осмотра"))
        self.records_tree.heading("organization", text="Организация ▼", command=lambda: self.open_records_column_filter("organization", "Организация"))
        self.records_tree.heading("name", text="ФИО ▼", command=lambda: self.open_records_column_filter("name", "ФИО"))
        self.records_tree.heading("birthday", text="Дата рождения ▼", command=lambda: self.open_records_column_filter("birthday", "Дата рождения"))
        self.records_tree.heading("sex", text="Пол ▼", command=lambda: self.open_records_column_filter("sex", "Пол"))
        self.records_tree.heading("division", text="Подразделение ▼", command=lambda: self.open_records_column_filter("division", "Подразделение"))
        self.records_tree.heading("profession", text="Должность ▼", command=lambda: self.open_records_column_filter("profession", "Должность"))
        self.records_tree.heading("factors", text="Факторы ▼", command=lambda: self.open_records_column_filter("factors", "Факторы"))
        self.records_tree.heading("typework", text="Виды работ ▼", command=lambda: self.open_records_column_filter("typework", "Виды работ"))
        self.records_tree.heading("diagnosis", text="Диагноз ▼", command=lambda: self.open_records_column_filter("diagnosis", "Диагноз"))
        self.records_tree.heading("card_number", text="№ карты ▼", command=lambda: self.open_records_column_filter("card_number", "№ карты"))
        self.records_tree.heading("ids_date", text="Дата ИДС ▼", command=lambda: self.open_records_column_filter("ids_date", "Дата ИДС"))
        self.records_tree.heading("date_created", text="Дата создания ▼", command=lambda: self.open_records_column_filter("date_created", "Дата создания"))

        self.records_tree.column("selected", width=44, anchor="center")
        self.records_tree.column("type", width=120, anchor="center")
        self.records_tree.column("organization", width=180, anchor="w")
        self.records_tree.column("name", width=220, anchor="w")
        self.records_tree.column("birthday", width=100, anchor="center")
        self.records_tree.column("sex", width=55, anchor="center")
        self.records_tree.column("division", width=150, anchor="w")
        self.records_tree.column("profession", width=150, anchor="w")
        self.records_tree.column("factors", width=160, anchor="w")
        self.records_tree.column("typework", width=160, anchor="w")
        self.records_tree.column("diagnosis", width=160, anchor="w")
        self.records_tree.column("card_number", width=90, anchor="center")
        self.records_tree.column("ids_date", width=100, anchor="center")
        self.records_tree.column("date_created", width=140, anchor="center")

        self.records_tree.tag_configure("odd", background="white")
        self.records_tree.tag_configure("even", background="#f3f3f3")

        self.records_tree.bind("<Button-1>", self.on_records_tree_click)
        self.records_tree.bind("<Double-1>", self.on_records_tree_double_click)
        self.records_search_var.trace_add("write", lambda *_: self.schedule_records_refresh())

        actions_frame = tk.Frame(top)
        actions_frame.pack(fill="x", padx=10, pady=(0, 10))

        tk.Button(
            actions_frame,
            text="Удалить выбранные записи",
            command=self.delete_selected_records,
            bg="#d9534f",
            fg="white"
        ).pack(side="right")

        tk.Button(
            actions_frame,
            text="Сформировать заключение",
            command=self.generate_selected_records_document,
            bg="#4CAF50",
            fg="white"
        ).pack(side="right", padx=(0, 8))


        self.refresh_records_tree()

    def reset_records_filters(self):
        self.records_date_from_var.set("")
        self.records_date_to_var.set("")
        self.records_search_var.set("")
        self.records_column_filters = {}
        self.refresh_records_tree()

    def close_records_filter_windows(self):
        alive_windows = []
        for win in self.records_filter_windows:
            try:
                if win and win.winfo_exists():
                    win.destroy()
            except Exception:
                pass
        self.records_filter_windows = alive_windows

    def apply_records_filters(self):
        self.close_records_filter_windows()
        self.refresh_records_tree()

    def schedule_records_refresh(self, delay=250):
        if self.records_search_after_id:
            try:
                self.after_cancel(self.records_search_after_id)
            except Exception:
                pass
        self.records_search_after_id = self.after(delay, self._run_scheduled_records_refresh)

    def _run_scheduled_records_refresh(self):
        self.records_search_after_id = None
        self.refresh_records_tree()

    def _records_column_title(self, column_name):
        titles = {
            "type": "Тип осмотра",
            "organization": "Организация",
            "name": "ФИО",
            "birthday": "Дата рождения",
            "sex": "Пол",
            "division": "Подразделение",
            "profession": "Должность",
            "factors": "Факторы",
            "typework": "Виды работ",
            "diagnosis": "Диагноз",
            "card_number": "№ карты",
            "ids_date": "Дата ИДС",
            "date_created": "Дата создания",
        }
        return titles.get(column_name, column_name)

    def _update_records_filter_headings(self):
        if not hasattr(self, "records_tree") or self.records_tree is None:
            return

        for column_name in (
            "type", "organization", "name", "birthday", "sex", "division", "profession",
            "factors", "typework", "diagnosis", "card_number", "ids_date", "date_created"
        ):
            title = self._records_column_title(column_name)
            has_filter = bool(self.records_column_filters.get(column_name))
            suffix = " ▼*" if has_filter else " ▼"
            self.records_tree.heading(
                column_name,
                text=f"{title}{suffix}",
                command=lambda c=column_name, t=title: self.open_records_column_filter(c, t)
            )

    def open_records_column_filter(self, column_name, title):
        all_records = self._read_all_xml_records()
        values = sorted({(record.get(column_name) or "").strip() for record in all_records}, key=lambda x: x.lower())

        top = tk.Toplevel(self)
        top.title(f"Фильтр: {title}")
        top.geometry("420x520")
        top.resizable(False, False)
        self.records_filter_windows.append(top)

        def on_filter_close():
            try:
                if top in self.records_filter_windows:
                    self.records_filter_windows.remove(top)
            except Exception:
                pass
            top.destroy()

        top.protocol("WM_DELETE_WINDOW", on_filter_close)

        tk.Label(top, text=f"Фильтр по столбцу: {title}", font=("Arial", 10, "bold")).pack(anchor="w", padx=10, pady=(10, 6))

        search_var = tk.StringVar()
        tk.Entry(top, textvariable=search_var).pack(fill="x", padx=10, pady=(0, 8))

        current_selected = set(self.records_column_filters.get(column_name, set(values))) if values else set()
        vars_map = {}

        controls_frame = tk.Frame(top)
        controls_frame.pack(fill="x", padx=10, pady=(0, 8))

        list_outer = tk.Frame(top)
        list_outer.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        canvas = tk.Canvas(list_outer, highlightthickness=0)
        scrollbar = ttk.Scrollbar(list_outer, orient="vertical", command=canvas.yview)
        inner = tk.Frame(canvas)

        inner.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=inner, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        def rebuild_list(*_):
            for child in inner.winfo_children():
                child.destroy()

            filter_text = search_var.get().strip().lower()
            visible_values = [v for v in values if filter_text in v.lower()]

            for value in visible_values:
                if value not in vars_map:
                    vars_map[value] = tk.BooleanVar(value=value in current_selected)
                cb = tk.Checkbutton(inner, text=value if value else "(пусто)", variable=vars_map[value], anchor="w", justify="left")
                cb.pack(fill="x", anchor="w")

        def select_all_visible():
            filter_text = search_var.get().strip().lower()
            for value in values:
                if filter_text in value.lower():
                    if value not in vars_map:
                        vars_map[value] = tk.BooleanVar(value=False)
                    vars_map[value].set(True)

        def clear_all_visible():
            filter_text = search_var.get().strip().lower()
            for value in values:
                if filter_text in value.lower():
                    if value not in vars_map:
                        vars_map[value] = tk.BooleanVar(value=False)
                    vars_map[value].set(False)

        tk.Button(controls_frame, text="Выделить все", command=lambda: (select_all_visible(), rebuild_list())).pack(side="left")
        tk.Button(controls_frame, text="Снять все", command=lambda: (clear_all_visible(), rebuild_list())).pack(side="left", padx=(6, 0))

        def apply_filter():
            selected_values = set()
            for value in values:
                if value not in vars_map:
                    vars_map[value] = tk.BooleanVar(value=value in current_selected)
                if vars_map[value].get():
                    selected_values.add(value)

            if len(selected_values) == len(values):
                self.records_column_filters.pop(column_name, None)
            else:
                self.records_column_filters[column_name] = selected_values

            self.refresh_records_tree()
            on_filter_close()

        def reset_filter():
            self.records_column_filters.pop(column_name, None)
            self.refresh_records_tree()
            on_filter_close()

        search_var.trace_add("write", rebuild_list)
        rebuild_list()

        buttons = tk.Frame(top)
        buttons.pack(fill="x", padx=10, pady=(0, 10))
        tk.Button(buttons, text="Сбросить фильтр", command=reset_filter).pack(side="left")
        tk.Button(buttons, text="Применить", command=apply_filter, bg="#4CAF50", fg="white").pack(side="right")

    def _parse_records_filter_date(self, value):
        value = (value or "").strip()
        if not value:
            return None
        try:
            return datetime.datetime.strptime(value, "%d.%m.%Y").date()
        except ValueError:
            return None

    def _read_all_xml_records(self, force_reload=False):
        if self.records_cache is not None and not force_reload:
            return [dict(record) for record in self.records_cache]

        if not os.path.exists(XML_PATH):
            self.records_cache = []
            return []

        tree = ET.parse(XML_PATH)
        root = tree.getroot()
        records = []
        for xml_index, person in enumerate(root.findall("person")):
            saved_id = person.findtext("id", default="")
            record = {
                "type": person.findtext("type", default="предварительный"),
                "organization": person.findtext("organization", default=""),
                "name": person.findtext("name", default=""),
                "birthday": person.findtext("birthday", default=""),
                "sex": person.findtext("sex", default=""),
                "division": person.findtext("division", default=""),
                "profession": person.findtext("profession", default=""),
                "factors": person.findtext("factors", default=""),
                "typework": person.findtext("typework", default=""),
                "diagnosis": person.findtext("diagnosis", default=""),
                "card_number": person.findtext("card_number", default=""),
                "ids_date": person.findtext("ids_date", default=""),
                "date_created": person.findtext("date_created", default=""),
                "id": saved_id,
                "_xml_index": xml_index,
            }
            record["_select_key"] = f"xml_index:{xml_index}|id:{saved_id}"
            records.append(record)
        self.records_cache = [dict(record) for record in records]
        return [dict(record) for record in self.records_cache]

    def refresh_records_tree(self):
        if not hasattr(self, "records_tree") or self.records_tree is None:
            return

        for item in self.records_tree.get_children():
            self.records_tree.delete(item)

        date_from = self._parse_records_filter_date(getattr(self, "records_date_from_var", tk.StringVar()).get())
        date_to = self._parse_records_filter_date(getattr(self, "records_date_to_var", tk.StringVar()).get())
        fio_search = getattr(self, "records_search_var", tk.StringVar()).get().strip().lower()

        records = self._read_all_xml_records()

        def ids_date_key(record):
            raw = (record.get("ids_date") or "").strip()
            try:
                return datetime.datetime.strptime(raw, "%d.%m.%Y")
            except ValueError:
                return datetime.datetime.min

        filtered = []
        for record in records:
            name = (record.get("name") or "").strip().lower()
            if fio_search and fio_search not in name:
                continue

            ids_date_raw = (record.get("ids_date") or "").strip()
            ids_date_obj = None
            if ids_date_raw:
                try:
                    ids_date_obj = datetime.datetime.strptime(ids_date_raw, "%d.%m.%Y").date()
                except ValueError:
                    ids_date_obj = None

            if date_from is not None:
                if ids_date_obj is None or ids_date_obj < date_from:
                    continue

            if date_to is not None:
                if ids_date_obj is None or ids_date_obj > date_to:
                    continue

            skip_record = False
            for column_name, allowed_values in self.records_column_filters.items():
                record_value = (record.get(column_name) or "").strip()
                if record_value not in allowed_values:
                    skip_record = True
                    break
            if skip_record:
                continue

            filtered.append(record)

        filtered.sort(key=ids_date_key, reverse=True)
        self.records_current_rows = filtered

        visible_keys = [record.get("_select_key") for record in filtered]
        self.records_all_visible_selected = bool(visible_keys) and all(key in self.records_selected_keys for key in visible_keys)
        self.records_tree.heading("selected", text="☑" if self.records_all_visible_selected else "☐", command=self.toggle_all_records_selection)
        self._update_records_filter_headings()

        for index, record in enumerate(filtered):
            tag = "odd" if index % 2 == 0 else "even"
            is_selected = record.get("_select_key") in self.records_selected_keys
            self.records_tree.insert(
                "",
                "end",
                iid=str(record.get("_select_key")),
                values=(
                    "☑" if is_selected else "☐",
                    record.get("type", "предварительный"),
                    record.get("organization", ""),
                    record.get("name", ""),
                    record.get("birthday", ""),
                    record.get("sex", ""),
                    record.get("division", ""),
                    record.get("profession", ""),
                    record.get("factors", ""),
                    record.get("typework", ""),
                    record.get("diagnosis", ""),
                    record.get("card_number", ""),
                    record.get("ids_date", ""),
                    record.get("date_created", ""),
                ),
                tags=(tag,)
            )

    def toggle_all_records_selection(self):
        if not hasattr(self, "records_current_rows"):
            return

        visible_keys = [record.get("_select_key") for record in self.records_current_rows]
        if not visible_keys:
            return

        all_selected = all(key in self.records_selected_keys for key in visible_keys)
        if all_selected:
            for key in visible_keys:
                self.records_selected_keys.discard(key)
        else:
            for key in visible_keys:
                self.records_selected_keys.add(key)

        self.refresh_records_tree()

    def on_records_tree_click(self, event):
        if not hasattr(self, "records_tree") or self.records_tree is None:
            return

        region = self.records_tree.identify("region", event.x, event.y)
        if region != "cell":
            return

        column = self.records_tree.identify_column(event.x)
        item_id = self.records_tree.identify_row(event.y)

        if column != "#1" or not item_id:
            return

        if item_id in self.records_selected_keys:
            self.records_selected_keys.discard(item_id)
        else:
            self.records_selected_keys.add(item_id)

        self.refresh_records_tree()
        return "break"

    def _get_selected_records(self):
        records = self._read_all_xml_records()
        return [record for record in records if record.get("_select_key") in self.records_selected_keys]

    def delete_selected_records(self):
        selected_records = self._get_selected_records()
        if not selected_records:
            messagebox.showerror("Ошибка", "Не выбрано ни одной записи")
            return

        confirm = messagebox.askyesno(
            "Подтверждение удаления",
            f"Удалить выбранные записи: {len(selected_records)} шт.?"
        )
        if not confirm:
            return

        if not os.path.exists(XML_PATH):
            messagebox.showerror("Ошибка", "Файл data.xml не найден")
            return

        tree = ET.parse(XML_PATH)
        root = tree.getroot()
        selected_indexes = {record.get("_xml_index") for record in selected_records}
        persons = root.findall("person")

        for idx in sorted(selected_indexes, reverse=True):
            if idx is None:
                continue
            if 0 <= idx < len(persons):
                root.remove(persons[idx])

        tree.write(XML_PATH, encoding="utf-8", xml_declaration=True)
        self.prettify_xml(XML_PATH)

        self.records_selected_keys.clear()
        self.records_cache = self._read_all_xml_records(force_reload=True)
        self.data = self.load_data()
        self.update_comboboxes()
        self.refresh_records_tree()
        self.show_notification(f"Удалено записей: {len(selected_records)}")
        messagebox.showinfo("Готово", f"Удалено записей: {len(selected_records)}")

    def generate_selected_records_document(self):
        selected_records = self._get_selected_records()
        if not selected_records:
            messagebox.showerror("Ошибка", "Не выбрано ни одной записи")
            return

        save_dir = self.settings.get("save_dir", os.getcwd())
        output_path = os.path.join(
            save_dir,
            f"заключения_из_выбранных_записей_{datetime.datetime.now().strftime('%d.%m.%Y_%H-%M-%S')}.docx"
        )

        temp_paths = []
        try:
            for record in selected_records:
                type_raw = record.get("type", "предварительный")
                type_genitive = {
                    "предварительный": "предварительного",
                    "периодический": "периодического"
                }.get(type_raw, type_raw)

                form_data = {
                    "{type}": type_genitive,
                    "{organization}": record.get("organization", ""),
                    "{name}": record.get("name", ""),
                    "{birthday}": record.get("birthday", ""),
                    "{sex}": record.get("sex", ""),
                    "{division}": record.get("division", ""),
                    "{profession}": record.get("profession", ""),
                    "{factors}": record.get("factors", ""),
                    "{typework}": record.get("typework", ""),
                    "{ids_date}": record.get("ids_date", ""),
                    "{diagnosis}": record.get("diagnosis", ""),
                    "{card_number}": record.get("card_number", ""),
                    "{year}": str(datetime.datetime.now().year)
                }

                temp_doc = Document(TEMPLATE_PATH)
                self.replace_placeholders(temp_doc, form_data)
                temp_doc_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
                temp_doc.save(temp_doc_path)
                temp_paths.append(temp_doc_path)

            if not temp_paths:
                messagebox.showerror("Ошибка", "Не удалось сформировать документы")
                return

            shutil.copyfile(temp_paths[0], output_path)
            if len(temp_paths) > 1:
                for temp_path in temp_paths[1:]:
                    self.append_doc_with_page_break(output_path, temp_path)

            self.show_notification(f"Файл сохранён: {output_path}")
            messagebox.showinfo("Готово", f"Сформирован файл:\n{output_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сформировать файл:\n{e}")
        finally:
            for temp_path in temp_paths:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass
    def on_records_tree_double_click(self, event):
        if not hasattr(self, "records_tree") or self.records_tree is None:
            return

        region = self.records_tree.identify("region", event.x, event.y)
        if region != "cell":
            return

        column = self.records_tree.identify_column(event.x)
        item_id = self.records_tree.identify_row(event.y)

        if not item_id:
            return

        # двойной клик по колонке чекбокса не открывает редактор
        if column == "#1":
            return

        self.open_record_edit_window(item_id)


    def open_record_edit_window(self, select_key):
        all_records = self._read_all_xml_records()
        record = next((r for r in all_records if r.get("_select_key") == select_key), None)
        if not record:
            messagebox.showerror("Ошибка", "Запись не найдена")
            return

        top = tk.Toplevel(self)
        top.title("Редактировать запись")
        top.geometry("720x520")
        top.resizable(False, False)

        fields = [
            ("type", "Тип осмотра"),
            ("organization", "Организация"),
            ("name", "ФИО"),
            ("birthday", "Дата рождения"),
            ("sex", "Пол"),
            ("division", "Подразделение"),
            ("profession", "Должность"),
            ("factors", "Факторы"),
            ("typework", "Виды работ"),
            ("diagnosis", "Диагноз"),
            ("card_number", "№ карты"),
            ("ids_date", "Дата ИДС"),
        ]

        vars_map = {}
        for row, (field_name, label_text) in enumerate(fields):
            tk.Label(top, text=label_text).grid(row=row, column=0, sticky="w", padx=10, pady=(10 if row == 0 else 6, 0))
            var = tk.StringVar(value=record.get(field_name, ""))
            vars_map[field_name] = var

            if field_name == "type":
                widget = Combobox(top, textvariable=var, values=["предварительный", "периодический"], width=47, state="readonly")
                widget.grid(row=row, column=1, sticky="ew", padx=10, pady=(10 if row == 0 else 6, 0))
            elif field_name == "sex":
                widget = Combobox(top, textvariable=var, values=["М", "Ж"], width=47, state="readonly")
                widget.grid(row=row, column=1, sticky="ew", padx=10, pady=(10 if row == 0 else 6, 0))
            else:
                widget = tk.Entry(top, textvariable=var, width=50)
                widget.grid(row=row, column=1, sticky="ew", padx=10, pady=(10 if row == 0 else 6, 0))
                if field_name in ("birthday", "ids_date"):
                    widget.bind("<KeyRelease>", self.format_date)

        top.grid_columnconfigure(1, weight=1)

        def save_changes():
            birthday = vars_map["birthday"].get().strip()
            ids_date = vars_map["ids_date"].get().strip()

            if not vars_map["organization"].get().strip():
                messagebox.showerror("Ошибка", "Организация не может быть пустой")
                return
            if not vars_map["name"].get().strip():
                messagebox.showerror("Ошибка", "ФИО не может быть пустым")
                return
            if birthday and not self.is_valid_date(birthday):
                messagebox.showerror("Ошибка", "Дата рождения должна быть в формате ДД.ММ.ГГГГ")
                return
            if ids_date and not self.is_valid_date(ids_date):
                messagebox.showerror("Ошибка", "Дата ИДС должна быть в формате ДД.ММ.ГГГГ")
                return

            if not os.path.exists(XML_PATH):
                messagebox.showerror("Ошибка", "Файл data.xml не найден")
                return

            tree = ET.parse(XML_PATH)
            root = tree.getroot()
            persons = root.findall("person")
            xml_index = record.get("_xml_index")
            if xml_index is None or xml_index < 0 or xml_index >= len(persons):
                messagebox.showerror("Ошибка", "Не удалось определить запись в XML")
                return

            person = persons[xml_index]
            for field_name, _label in fields:
                elem = person.find(field_name)
                if elem is None:
                    elem = ET.SubElement(person, field_name)
                elem.text = vars_map[field_name].get().strip()

            tree.write(XML_PATH, encoding="utf-8", xml_declaration=True)
            self.prettify_xml(XML_PATH)
            top.destroy()
            self.records_cache = self._read_all_xml_records(force_reload=True)
            self.data = self.load_data()
            self.update_comboboxes()
            self.refresh_records_tree()
            self.show_notification("Запись обновлена")

        button_row = tk.Frame(top)
        button_row.grid(row=len(fields), column=0, columnspan=2, sticky="ew", padx=10, pady=16)
        tk.Button(button_row, text="Сохранить", command=save_changes, bg="#4CAF50", fg="white").pack(side="right")
        tk.Button(button_row, text="Отмена", command=top.destroy).pack(side="right", padx=(0, 8))